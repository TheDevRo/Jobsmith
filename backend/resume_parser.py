"""
resume_parser.py — Extract structured profile data from a résumé.

Two responsibilities:

  1. extract_text()  — pull plain text out of an uploaded PDF / DOCX / TXT.
  2. parse_resume()  — ask the local LLM to map that text onto the
     UserProfile schema, strictly EXTRACTIVELY (never invent data).

Used only by the first-run onboarding wizard. The result is shown to the
user for review/edit before anything is persisted — this module never
writes config.
"""

from __future__ import annotations

import io
import json
import logging
import re

from . import ai_engine

logger = logging.getLogger(__name__)

# Fields the wizard can prefill from a résumé. Demographic / credential /
# salary fields are intentionally excluded — they are not on a résumé and
# must be entered deliberately by the user.
_STR_FIELDS = (
    "full_name", "email", "phone", "location",
    "street_address", "street_address_2", "city", "state", "zip_code",
    "linkedin", "github", "portfolio", "summary",
)
_MAX_CHARS = 16000  # keep the prompt inside the local model's context window


# ---------------------------------------------------------------------------
# Text extraction
# ---------------------------------------------------------------------------
def extract_text(filename: str, data: bytes) -> str:
    """Return plain text from an uploaded résumé.

    Supports .pdf (pypdf), .docx (python-docx) and .txt / plain text.
    Raises ValueError for unsupported or unreadable files.
    """
    name = (filename or "").lower().strip()

    if name.endswith(".pdf"):
        try:
            from pypdf import PdfReader
        except ImportError as exc:  # pragma: no cover - dependency guard
            raise ValueError("PDF support requires the 'pypdf' package") from exc
        try:
            reader = PdfReader(io.BytesIO(data))
            pages = [(page.extract_text() or "") for page in reader.pages]
        except Exception as exc:
            raise ValueError(f"Could not read PDF: {exc}") from exc
        return "\n".join(pages).strip()

    if name.endswith(".docx"):
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - dependency guard
            raise ValueError("DOCX support requires the 'python-docx' package") from exc
        try:
            doc = Document(io.BytesIO(data))
        except Exception as exc:
            raise ValueError(f"Could not read DOCX: {exc}") from exc
        lines = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append("  ".join(cells))
        return "\n".join(lines).strip()

    if name.endswith((".txt", ".md", ".text")) or not name:
        try:
            return data.decode("utf-8", errors="replace").strip()
        except Exception as exc:
            raise ValueError(f"Could not decode text file: {exc}") from exc

    raise ValueError(
        f"Unsupported file type: {filename!r}. Upload a PDF, DOCX, or TXT, "
        "or paste the résumé text instead."
    )


# ---------------------------------------------------------------------------
# LLM extraction
# ---------------------------------------------------------------------------
_PROMPT = """You are a résumé parser. Extract ONLY information that is \
literally present in the résumé text below. Do NOT infer, guess, embellish, \
or invent anything. If a field is not clearly stated in the résumé, return an \
empty string "" (or an empty list for list fields). Never fabricate names, \
employers, dates, contact details, schools, or skills.

Return ONLY a single JSON object, no prose, no markdown fences, with EXACTLY \
these keys:

{{
  "full_name": "",
  "email": "",
  "phone": "",
  "location": "",
  "street_address": "",
  "street_address_2": "",
  "city": "",
  "state": "",
  "zip_code": "",
  "linkedin": "",
  "github": "",
  "portfolio": "",
  "summary": "",
  "skills": [],
  "experience": [
    {{"title": "", "company": "", "start_date": "", "end_date": "Present", "bullets": []}}
  ],
  "education": [
    {{"degree": "", "school": "", "year": ""}}
  ],
  "certifications": []
}}

Rules:
- "location" should be "City, ST" if present; also fill city/state/zip_code \
when an explicit address is given.
- Dates: keep them as written in the résumé (e.g. "2021", "Jan 2021", \
"2021-01"). Use "Present" for a current role's end_date.
- "bullets": copy the résumé's accomplishment lines verbatim, lightly \
trimmed; do not rewrite them.
- "skills": only skills explicitly listed; one skill per array item.
- Omit empty experience/education objects entirely rather than padding.

RÉSUMÉ TEXT:
\"\"\"
{resume}
\"\"\"

Return only the JSON object."""


def _coerce_str_list(value) -> list[str]:
    if isinstance(value, list):
        return [str(v).strip() for v in value if str(v).strip()]
    if isinstance(value, str) and value.strip():
        return [s.strip() for s in re.split(r"[,\n;]", value) if s.strip()]
    return []


def _sanitize(raw: dict) -> dict:
    """Coerce the model's JSON onto the UserProfile partial shape.

    Drops unknown keys, fixes types, and removes empty experience/education
    rows so the review screen isn't littered with blanks.
    """
    out: dict = {}
    for f in _STR_FIELDS:
        v = raw.get(f, "")
        out[f] = v.strip() if isinstance(v, str) else ("" if v is None else str(v))

    out["skills"] = _coerce_str_list(raw.get("skills"))
    out["certifications"] = _coerce_str_list(raw.get("certifications"))

    experience = []
    for e in raw.get("experience") or []:
        if not isinstance(e, dict):
            continue
        title = str(e.get("title", "")).strip()
        company = str(e.get("company", "")).strip()
        if not title and not company:
            continue
        experience.append({
            "title": title,
            "company": company,
            "start_date": str(e.get("start_date", "")).strip(),
            "end_date": str(e.get("end_date", "") or "Present").strip(),
            "bullets": _coerce_str_list(e.get("bullets")),
        })
    out["experience"] = experience

    education = []
    for e in raw.get("education") or []:
        if not isinstance(e, dict):
            continue
        degree = str(e.get("degree", "")).strip()
        school = str(e.get("school", "")).strip()
        if not degree and not school:
            continue
        education.append({
            "degree": degree,
            "school": school,
            "year": str(e.get("year", "")).strip(),
        })
    out["education"] = education
    return out


def _extract_json(text: str) -> dict:
    """Best-effort JSON recovery, mirroring ai_engine.score_job_fit fallbacks."""
    text = text.strip()
    # Strip ```json fences if the model added them
    fenced = re.search(r"```(?:json)?\s*(\{.*\})\s*```", text, re.DOTALL)
    if fenced:
        text = fenced.group(1)
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if match:
        try:
            return json.loads(match.group())
        except json.JSONDecodeError:
            pass
    raise ValueError("Model did not return parseable JSON")


async def parse_resume(text: str, config: dict, prompt_template: str | None = None) -> dict:
    """Extract a partial profile dict from résumé-like text via the local LLM.

    `prompt_template` lets other extractive sources (e.g. the LinkedIn
    profile importer) supply their own prompt; it must contain a `{resume}`
    placeholder and request the same JSON schema as `_PROMPT`.

    Returns {"profile": {...}, "warnings": [...]}. Never raises for a bad
    model response — instead returns an empty profile plus a warning so the
    user can still fill the form manually.
    """
    warnings: list[str] = []
    text = (text or "").strip()
    if not text:
        return {"profile": _sanitize({}), "warnings": ["No résumé text to parse."]}

    if len(text) > _MAX_CHARS:
        text = text[:_MAX_CHARS]
        warnings.append(
            "The text was long; only the first part was parsed. Review fields carefully."
        )

    ai_cfg = config.get("ai", {})
    client = ai_engine._get_client(config, "strong")
    prompt = (prompt_template or _PROMPT).format(resume=text)

    try:
        response = await client.chat.completions.create(
            model=ai_engine._model(config, "strong"),
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=ai_cfg.get("max_tokens", 4096),
        )
        raw_text = (response.choices[0].message.content or "").strip()
    except Exception as exc:
        logger.exception("Résumé parse: LLM call failed")
        return {
            "profile": _sanitize({}),
            "warnings": [f"AI extraction failed ({exc}). Fill the form manually."],
        }

    try:
        data = _extract_json(raw_text)
    except ValueError:
        logger.warning("Résumé parse: unparseable model output: %s", raw_text[:300])
        return {
            "profile": _sanitize({}),
            "warnings": [
                "Could not extract structured data automatically. Fill the form "
                "manually or try again."
            ],
        }

    profile = _sanitize(data if isinstance(data, dict) else {})
    if not profile.get("full_name") and not profile.get("experience"):
        warnings.append(
            "Little structured data was found — double-check every field below."
        )
    return {"profile": profile, "warnings": warnings}
