"""
resume_generator.py — Generate professional DOCX resume and cover letter files.

Uses python-docx to create clean, ATS-friendly formatted documents.
"""

import logging
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from .paths import project_root

logger = logging.getLogger(__name__)

RESUMES_DIR = project_root() / "resumes"

# -- Color palette (fixed body-text colors shared by every style) --
COLOR_BLACK = RGBColor(0x22, 0x22, 0x22)
COLOR_DARK = RGBColor(0x33, 0x33, 0x33)
COLOR_GRAY = RGBColor(0x66, 0x66, 0x66)


def _rgb(hex6: str) -> RGBColor:
    """Convert a 6-char hex string ('1F3A5F') to a python-docx RGBColor."""
    return RGBColor.from_string(hex6)


# User-selectable accent palette. "default" keeps each preset's own accent.
ACCENT_CHOICES: dict[str, str] = {
    "navy": "1F3A5F",
    "burgundy": "6D1F2C",
    "forest": "1F4D3A",
    "plum": "3D3A4F",
    "charcoal": "37404A",
}

DEFAULT_RESUME_STYLE = "ledger"

# Configs written before the 5-style lineup carry the retired names.
LEGACY_STYLE_ALIASES: dict[str, str] = {
    "standard": "ledger",
    "modern": "ledger",
    "minimal": "swiss",
}

# Tokens every preset inherits; presets override only what differs.
# All colors are 6-char hex strings; the value "accent" resolves to the
# active accent (the preset's own, or the user's resume_accent choice).
_STYLE_DEFAULTS: dict = {
    "body_font": "Calibri",
    "name_font": "Calibri",
    "body_size": 10.5,
    "name_size": 20,
    "name_bold": True,
    "name_uppercase": False,
    "name_small_caps": False,
    "name_letter_spacing": 0,
    "name_align": "left",
    "name_color": "222222",
    "accent": "1F3A5F",       # preset's default accent hex
    "accent_locked": False,   # True → deliberately monochrome; ignore user accent
    "accent_color": "accent",
    "contact_color": "666666",
    "contact_separator": "  ·  ",
    "contact_inline": False,  # contact info on the name line (compact)
    "header_size": 10,
    "header_color": "222222",
    "header_small_caps": False,
    "header_letter_spacing": 0,
    "header_underline": False,
    "header_underline_color": "999999",
    "header_underline_size": "4",
    "header_rule_style": "full",   # full-width rule | short "stub" bar
    "section_gap": 10,             # points of space before section headers
    "bullet_marker": "•  ",
    "bullet_marker_size": 10.5,
    "bullet_marker_color": "333333",
    "name_rule": False,
    "name_rule_style": "single",   # single | double | stub
    "name_rule_color": "999999",
    "name_rule_size": "6",
    "banner": False,               # shaded band behind name + contact
    "margins": (0.6, 0.6, 0.75, 0.75),  # top, bottom, left, right (inches)
    "line_spacing": None,
    "entry_layout": "inline",      # "Title · Company ....... Dates" one line
    "entry_separator": "  ·  ",
    "company_style": "plain",      # plain | italic | accent
    "skills_separator": ", ",
    "hyperlinks": True,
}

# Style tokens — each preset is a flat dict consumed by generate_resume_docx.
# All presets stay single-column, real text, no images/tables (ATS-safe).
_STYLES: dict[str, dict] = {
    # Executive — engraved-letterhead serif: centered small-caps name over a
    # thin double rule. Deliberately monochrome (accent_locked).
    "executive": {
        "body_font": "Georgia",
        "name_font": "Georgia",
        "name_size": 22,
        "name_bold": False,
        "name_small_caps": True,
        "name_letter_spacing": 3,
        "name_align": "center",
        "name_color": "17202B",
        "accent": "17202B",
        "accent_locked": True,
        "header_size": 11,
        "header_color": "17202B",
        "header_small_caps": True,
        "header_letter_spacing": 1.5,
        "header_underline": True,
        "header_underline_color": "C9C2B4",
        "header_underline_size": "4",
        "name_rule": True,
        "name_rule_style": "double",
        "name_rule_color": "17202B",
        "name_rule_size": "4",
        "margins": (0.7, 0.7, 0.85, 0.85),
        "entry_separator": ", ",
        "company_style": "italic",
        "skills_separator": "  ·  ",
        "bullet_marker": "•  ",
        "bullet_marker_size": 9,
        "bullet_marker_color": "17202B",
    },
    # Ledger — bold sans with a short thick accent stub under the name and
    # stub-underlined section headers. The default style; shows off accents.
    #
    # Calibri, not Aptos: Aptos ships only with Microsoft 365, and every other
    # reader (LibreOffice, older Word, Google Docs) substitutes a *serif* for
    # it — which wrecks a style whose whole identity is "bold sans". The PDF
    # path embeds Lato, so PDFs keep the more modern face regardless.
    "ledger": {
        "body_font": "Calibri",
        "name_font": "Calibri",
        "name_size": 26,
        "name_color": "171C24",
        "header_size": 10,
        "header_color": "accent",
        "header_letter_spacing": 1.2,
        "header_underline": True,
        "header_underline_color": "accent",
        "header_underline_size": "20",
        "header_rule_style": "stub",
        "name_rule": True,
        "name_rule_style": "stub",
        "name_rule_color": "accent",
        "name_rule_size": "28",
        "margins": (0.6, 0.6, 0.8, 0.8),
        "line_spacing": 1.12,
        "company_style": "accent",
        "skills_separator": "  ·  ",
    },
    # Banner — full-width ink band behind name + contact. Paragraph shading
    # on real text (w:shd / backColor) — parses identically to plain text.
    # Calibri for the same reason as Ledger (see above).
    "banner": {
        "body_font": "Calibri",
        "name_font": "Calibri",
        "accent": "1F2D42",
        "name_size": 24,
        "name_color": "FFFFFF",
        "contact_color": "D7DEE9",
        "banner": True,
        "header_size": 10,
        "header_color": "accent",
        "header_letter_spacing": 1.2,
        "header_underline": True,
        "header_underline_color": "accent",
        "header_underline_size": "12",
        "margins": (0.5, 0.6, 0.8, 0.8),
        "line_spacing": 1.1,
        "entry_separator": "  —  ",
        "skills_separator": "  ·  ",
    },
    # Compact — 9.5pt, half-inch margins, contact on the name line,
    # pipe-separated skills. Two pages become one.
    "compact": {
        "body_size": 9.5,
        "name_size": 14,
        "contact_inline": True,
        "accent": "37404A",
        "header_size": 9,
        "header_color": "accent",
        "header_letter_spacing": 1,
        "header_underline": True,
        "header_underline_color": "CCCCCC",
        "header_underline_size": "4",
        "section_gap": 7,
        "name_rule": True,
        "name_rule_color": "999999",
        "name_rule_size": "4",
        "margins": (0.5, 0.5, 0.5, 0.5),
        "line_spacing": 1.0,
        "skills_separator": " | ",
        "bullet_marker_size": 9.5,
    },
    # Swiss — no rules, no color: hierarchy from spacing, weight, and quiet
    # letter-spaced grey headers alone. Deliberately monochrome.
    "swiss": {
        "body_font": "Arial",
        "name_font": "Arial",
        "name_size": 21,
        "name_bold": False,
        "name_color": "14181D",
        "accent": "3A3F47",
        "accent_locked": True,
        "header_size": 9.5,
        "header_color": "9AA0A9",
        "header_letter_spacing": 2.5,
        "section_gap": 15,
        "margins": (0.9, 0.8, 0.95, 0.95),
        "line_spacing": 1.2,
        "entry_separator": ", ",
        "skills_separator": "  ·  ",
        "bullet_marker": "—  ",
        "bullet_marker_color": "6E747D",
    },
}


def _resolve_resume_style(config: dict | None) -> dict:
    """Resolve config to a fully-populated style dict — shared by the DOCX
    and PDF paths for both resumes and cover letters.

    Handles legacy style names, merges preset over defaults, and substitutes
    the "accent" sentinel with the active accent hex (the user's
    application_honesty.resume_accent choice unless the preset is locked).
    """
    honesty = (config or {}).get("application_honesty", {})
    name = str(honesty.get("resume_style", DEFAULT_RESUME_STYLE)).lower()
    name = LEGACY_STYLE_ALIASES.get(name, name)
    preset = _STYLES.get(name, _STYLES[DEFAULT_RESUME_STYLE])

    s = dict(_STYLE_DEFAULTS)
    s.update(preset)

    accent_hex = s["accent"]
    if not s["accent_locked"]:
        choice = str(honesty.get("resume_accent", "default")).lower()
        accent_hex = ACCENT_CHOICES.get(choice, accent_hex)

    # Only *color* tokens carry the "accent" sentinel — company_style legitimately
    # has "accent" as an enum value and must not be rewritten to a hex.
    resolved = {
        k: (accent_hex if (v == "accent" and k.endswith("_color")) else v)
        for k, v in s.items()
    }
    resolved["accent"] = accent_hex
    return resolved


def _add_hyperlink(paragraph, url, text, font_name, size_pt, color_hex):
    """Insert a clickable hyperlink run into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))  # half-points
    rPr.append(sz)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), color_hex)
    rPr.append(color)

    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _ensure_dir():
    RESUMES_DIR.mkdir(parents=True, exist_ok=True)


def _set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def _set_paragraph_spacing(para, before=0, after=0, line_spacing=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = Pt(line_spacing)


def _add_bottom_border(paragraph, color="999999", size="6", val="single"):
    """Add a bottom border line to a paragraph. val="double" draws the thin
    twin rule the Executive style uses under its letterhead."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="{val}" w:sz="{size}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def _shade_paragraph(paragraph, fill_hex):
    """Fill a paragraph's background with a solid color (the Banner band).

    Shading is presentation-only in OOXML — the run text underneath is
    ordinary text, so ATS extraction is unaffected.
    """
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill_hex}"/>'
    ))


def _add_stub_rule(doc, content_width, width, color, size, before=2, after=4):
    """Add a short accent bar (the Ledger stub) as its own paragraph.

    A paragraph border spans the full text column, so a short bar is drawn by
    pushing the paragraph's right indent in until only `width` inches remain.
    The paragraph is collapsed to a hairline — an empty paragraph at body size
    would otherwise reserve a full blank line under every bar.
    """
    para = doc.add_paragraph()
    _set_paragraph_spacing(para, before=before, after=after)
    para.paragraph_format.right_indent = Inches(max(0.0, content_width - width))
    para.paragraph_format.line_spacing = Pt(1)  # exact — collapse the empty line
    _set_font(para.add_run(""), size=1)
    _add_bottom_border(para, color=color, size=size)
    return para


def _parse_resume_sections(content: str) -> dict:
    """
    Parse AI-generated resume text into sections.

    Case-insensitive matching against a broad set of header variants.
    If no recognised sections are detected, the entire content is treated
    as a summary so the DOCX is always generated rather than silently empty.
    """
    # Map every accepted variant → canonical section key
    _HEADER_MAP: dict[str, str] = {
        # Summary variants
        "SUMMARY": "summary",
        "PROFESSIONAL SUMMARY": "summary",
        "PROFILE SUMMARY": "summary",
        "PROFILE": "summary",
        "ABOUT ME": "summary",
        # Skills variants
        "SKILLS": "skills",
        "TECHNICAL SKILLS": "skills",
        "CORE COMPETENCIES": "skills",
        "KEY SKILLS": "skills",
        "COMPETENCIES": "skills",
        # Experience variants
        "EXPERIENCE": "experience",
        "PROFESSIONAL EXPERIENCE": "experience",
        "WORK EXPERIENCE": "experience",
        "EMPLOYMENT HISTORY": "experience",
        "CAREER HISTORY": "experience",
        # Education variants
        "EDUCATION": "education",
        "EDUCATIONAL BACKGROUND": "education",
        "ACADEMIC BACKGROUND": "education",
        # Certification variants
        "CERTIFICATIONS": "certifications",
        "CERTIFICATES": "certifications",
        "LICENSES": "certifications",
        "LICENSES & CERTIFICATIONS": "certifications",
        # Misc — kept for completeness but not rendered specially
        "PROJECTS": "projects",
        "AWARDS": "awards",
        "ACHIEVEMENTS": "awards",
    }

    sections: dict[str, str] = {}
    current_section = "preamble"
    current_lines: list[str] = []

    for line in content.split("\n"):
        stripped = line.strip()
        # Strip markdown prefix characters the AI might add despite instructions
        cleaned = re.sub(r'^[#*_`]+\s*', '', stripped).strip()
        # Normalise: upper-case and strip trailing colon/whitespace
        cleaned_upper = cleaned.upper().rstrip(": ")

        if cleaned_upper in _HEADER_MAP:
            if current_lines:
                sections[current_section] = "\n".join(current_lines).strip()
            current_section = _HEADER_MAP[cleaned_upper]
            current_lines = []
        else:
            current_lines.append(line)

    if current_lines:
        sections[current_section] = "\n".join(current_lines).strip()

    # Remove the preamble key — it's just lines before the first header
    sections.pop("preamble", None)

    # Fallback: if the LLM returned freeform text with no recognisable headers,
    # treat the whole content as a summary so we always produce a DOCX.
    if not sections:
        logger.warning(
            "_parse_resume_sections: no section headers found — treating full "
            "content (%d chars) as summary. Check LLM output format.", len(content)
        )
        sections["summary"] = content.strip()

    return sections


def _parse_experience_entries(text: str) -> list[dict]:
    """
    Parse experience section into structured entries.
    Each entry: {title, company, dates, bullets[]}

    Supported formats (in priority order):
      Structured  — Title: / Company: / Dates: prefix lines (from prompt)
      Pipe        — "Job Title | Company Name | Jan 2020 - Present"
      At-notation — "Job Title at Company Name (2020-2023)"
      Comma+paren — "Job Title, Company Name (2020-2023)"
      Fallback    — any non-bullet line starts a new entry; raw text preserved
    """
    entries = []
    current: dict | None = None

    def _save_current():
        if current:
            entries.append(current)

    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        # Strip markdown bold markers
        stripped = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)

        # ── Structured prefix format (highest priority) ──────────────────
        title_match   = re.match(r'^Title:\s*(.+)',   stripped, re.IGNORECASE)
        company_match = re.match(r'^Company:\s*(.+)', stripped, re.IGNORECASE)
        dates_match   = re.match(r'^Dates?:\s*(.+)',  stripped, re.IGNORECASE)

        if title_match:
            _save_current()
            current = {"title": title_match.group(1).strip(), "company": "", "dates": "", "bullets": []}
            continue

        if company_match and current:
            current["company"] = company_match.group(1).strip()
            continue

        if dates_match and current:
            current["dates"] = dates_match.group(1).strip()
            continue

        # ── Bullet points ─────────────────────────────────────────────────
        if stripped.startswith(("-", "•", "*", "–", "▸")):
            bullet = re.sub(r'^[-•*–▸]+\s*', '', stripped).strip()
            if bullet:
                if current is None:
                    # Orphan bullet before any entry — create a blank entry
                    current = {"title": "", "company": "", "dates": "", "bullets": []}
                current["bullets"].append(bullet)
            continue

        # ── Freeform header lines (no structured prefix) ─────────────────
        # Only try to parse as a header if it doesn't look like body text
        # (body text sentences are long; headers are short and concise).
        if len(stripped) < 120:
            # Format 1: "Title | Company | Dates"
            pipe_parts = [p.strip() for p in stripped.split("|")]
            if len(pipe_parts) >= 2:
                _save_current()
                current = {
                    "title":   pipe_parts[0],
                    "company": pipe_parts[1],
                    "dates":   pipe_parts[2] if len(pipe_parts) >= 3 else "",
                    "bullets": [],
                }
                continue

            # Format 2: "Title at Company (dates)" or "Title at Company"
            at_match = re.match(
                r'^(.+?)\s+at\s+(.+?)(?:\s+\(([^)]+)\))?$', stripped, re.IGNORECASE
            )
            if at_match:
                _save_current()
                current = {
                    "title":   at_match.group(1).strip(),
                    "company": at_match.group(2).strip(),
                    "dates":   at_match.group(3).strip() if at_match.group(3) else "",
                    "bullets": [],
                }
                continue

            # Format 3: "Title, Company (dates)"
            comma_paren = re.match(
                r'^(.+?),\s+(.+?)\s+\(([^)]+)\)$', stripped
            )
            if comma_paren:
                _save_current()
                current = {
                    "title":   comma_paren.group(1).strip(),
                    "company": comma_paren.group(2).strip(),
                    "dates":   comma_paren.group(3).strip(),
                    "bullets": [],
                }
                continue

        # ── Fallback: treat as a new entry title, preserving the raw text ─
        _save_current()
        current = {"title": stripped, "company": "", "dates": "", "bullets": []}

    _save_current()
    return entries


def _parse_education_entries(text: str) -> list[dict]:
    """
    Parse education section into structured entries: {degree, school, year}.

    Supported formats:
      Structured  — Degree: / School: / Year: prefix lines (from prompt)
      Freeform    — "B.S. Computer Science, Stanford University, 2020"
                    "B.S. Computer Science | Stanford University | 2020"
                    "B.S. Computer Science, Stanford University (2020)"
    """
    entries = []
    current: dict | None = None

    # Pattern to detect a 4-digit year
    _YEAR_RE = re.compile(r'\b(19|20)\d{2}\b')

    def _extract_year(s: str) -> tuple[str, str]:
        """Return (text_without_year, year_string). year is '' if not found."""
        m = _YEAR_RE.search(s)
        if m:
            year = m.group()
            rest = (s[:m.start()] + s[m.end():]).strip(" ,()–-")
            return rest, year
        return s, ""

    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue

        stripped = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)

        # ── Structured prefix lines ──────────────────────────────────────
        degree_match = re.match(r'^Degree:\s*(.+)', stripped, re.IGNORECASE)
        school_match = re.match(r'^School:\s*(.+)', stripped, re.IGNORECASE)
        year_match   = re.match(r'^Year:\s*(.+)',   stripped, re.IGNORECASE)

        if degree_match:
            if current:
                entries.append(current)
            current = {"degree": degree_match.group(1).strip(), "school": "", "year": ""}
            continue

        if school_match and current:
            current["school"] = school_match.group(1).strip()
            continue

        if year_match and current:
            current["year"] = year_match.group(1).strip()
            continue

        # Skip bullet points
        if stripped.startswith(("-", "•", "*", "–")):
            continue

        # ── Freeform line ─────────────────────────────────────────────────
        # Try pipe-separated first: "Degree | School | Year"
        pipe_parts = [p.strip() for p in stripped.split("|")]
        if len(pipe_parts) >= 2:
            if current:
                entries.append(current)
            rest, year = _extract_year(pipe_parts[-1])
            if year:
                school = pipe_parts[1].strip() if len(pipe_parts) >= 3 else rest
                current = {"degree": pipe_parts[0].strip(), "school": school, "year": year}
            else:
                current = {"degree": pipe_parts[0].strip(), "school": pipe_parts[1].strip(), "year": ""}
            continue

        # Try comma-separated: "Degree, School, Year" or "Degree, School (Year)"
        # Extract year first (could be in parens or as last comma token)
        paren_year = re.search(r'\((\d{4})\)', stripped)
        if paren_year:
            year = paren_year.group(1)
            without_year = stripped[:paren_year.start()].strip().rstrip(",")
            comma_parts = [p.strip() for p in without_year.split(",", 1)]
            if len(comma_parts) >= 2:
                if current:
                    entries.append(current)
                current = {"degree": comma_parts[0], "school": comma_parts[1], "year": year}
                continue

        comma_parts = [p.strip() for p in stripped.split(",")]
        if len(comma_parts) >= 2:
            last, year = _extract_year(comma_parts[-1])
            if year:
                school = ",".join(comma_parts[1:-1]).strip() or last
                if current:
                    entries.append(current)
                current = {"degree": comma_parts[0], "school": school, "year": year}
                continue

        # Last resort: treat the whole line as a degree field
        if current:
            entries.append(current)
        current = {"degree": stripped.lstrip("-•* "), "school": "", "year": ""}

    if current:
        entries.append(current)

    return entries


def _contact_entries(profile: dict, *, include_links: bool) -> list[tuple[str, str | None]]:
    """Contact bits as (display_text, url_or_None), in display order."""
    entries: list[tuple[str, str | None]] = []
    for field in ("email", "phone", "location"):
        val = profile.get(field, "")
        if val:
            entries.append((val, None))
    if not include_links:
        return entries
    for field in ("linkedin", "portfolio", "github"):
        url = profile.get(field, "")
        if not url:
            continue
        clean = re.sub(r"^https?://", "", url).rstrip("/")
        if clean:
            full_url = url if url.startswith("http") else f"https://{clean}"
            entries.append((clean, full_url))
    return entries


def _add_letterhead_docx(doc, profile, s, content_width, *, include_links=True):
    """Render the name + contact letterhead shared by resumes and cover letters.

    Honors name alignment, small caps, the Banner band, and the name rule
    (single / double / stub) so both documents in a pair look like one set.
    """
    body_font = s["body_font"]
    body_size = s["body_size"]
    contact_size = max(8.0, body_size - 1.5)
    align = (WD_ALIGN_PARAGRAPH.CENTER if s["name_align"] == "center"
             else WD_ALIGN_PARAGRAPH.LEFT)
    banner = s["banner"]

    def _shaded_spacer(height_pt):
        """Padding inside the band — paragraph shading doesn't cover the
        space before/after, so the band needs its own empty shaded lines."""
        pad = doc.add_paragraph()
        _set_paragraph_spacing(pad, before=0, after=0)
        pad.paragraph_format.line_spacing = 1.0
        _set_font(pad.add_run(""), name=body_font, size=height_pt)
        _shade_paragraph(pad, s["accent"])

    if banner:
        _shaded_spacer(5)

    name = profile.get("full_name", "")
    if name:
        name_para = doc.add_paragraph()
        name_para.alignment = align
        _set_paragraph_spacing(name_para, before=0, after=2)
        name_para.paragraph_format.line_spacing = 1.0
        display_name = name.upper() if s["name_uppercase"] else name
        name_run = name_para.add_run(display_name)
        _set_font(name_run, name=s["name_font"], size=s["name_size"],
                  bold=s["name_bold"], color=_rgb(s["name_color"]))
        if s["name_small_caps"]:
            name_run.font.small_caps = True
        if s["name_letter_spacing"]:
            name_run.font.letter_spacing = Pt(s["name_letter_spacing"])
        if banner:
            _shade_paragraph(name_para, s["accent"])

    entries = _contact_entries(profile, include_links=include_links)
    if entries:
        # Compact runs the contact line onto the name paragraph itself.
        inline = s["contact_inline"] and name
        if inline:
            contact_para = name_para
            contact_para.add_run("   ")
        else:
            contact_para = doc.add_paragraph()
            contact_para.alignment = align
            _set_paragraph_spacing(contact_para, before=0, after=2)
            contact_para.paragraph_format.line_spacing = 1.0

        link_hex = s["accent_color"]
        contact_color = _rgb(s["contact_color"])
        for i, (text, url) in enumerate(entries):
            if i > 0:
                sep_run = contact_para.add_run(s["contact_separator"])
                _set_font(sep_run, name=body_font, size=contact_size,
                          color=contact_color)
            if url and s["hyperlinks"] and not banner:
                _add_hyperlink(contact_para, url, text, body_font,
                               contact_size, link_hex)
            else:
                run = contact_para.add_run(text)
                _set_font(run, name=body_font, size=contact_size,
                          color=contact_color)
        if banner and not inline:
            _shade_paragraph(contact_para, s["accent"])

    if banner:
        _shaded_spacer(6)
        return

    if s["name_rule"]:
        rule_style = s["name_rule_style"]
        if rule_style == "stub":
            _add_stub_rule(doc, content_width, 0.55, s["name_rule_color"],
                           s["name_rule_size"], before=3, after=6)
        else:
            rule_para = doc.add_paragraph()
            _set_paragraph_spacing(rule_para, before=4, after=4)
            _add_bottom_border(
                rule_para, color=s["name_rule_color"], size=s["name_rule_size"],
                val="double" if rule_style == "double" else "single",
            )


def generate_resume_docx(content: str, profile: dict, job: dict, config: dict | None = None) -> str:
    """
    Create a professional, ATS-friendly resume DOCX from AI-generated content.

    config.application_honesty.resume_style picks the visual theme:
      "executive" — Georgia serif, centered small-caps name, double rule
      "ledger"    — bold sans, accent stub bar, accent company names (default)
      "banner"    — shaded ink band behind the name block
      "compact"   — 9.5pt, tight margins, contact on the name line
      "swiss"     — no rules, no color; spacing and weight only

    config.application_honesty.resume_accent recolors the accent-driven styles.
    Every style stays single-column real text — no images, tables or text boxes.
    Returns the file path to the generated document.
    """
    _ensure_dir()

    s = _resolve_resume_style(config)

    _font_name   = s["body_font"]
    _body_size   = s["body_size"]
    _header_size = s["header_size"]
    _accent_hex  = s["accent_color"]
    _dates_size  = max(8.0, _body_size - 0.5)

    doc = Document()

    # -- Page margins --
    top_m, bot_m, left_m, right_m = s["margins"]
    for section in doc.sections:
        section.top_margin    = Inches(top_m)
        section.bottom_margin = Inches(bot_m)
        section.left_margin   = Inches(left_m)
        section.right_margin  = Inches(right_m)

    content_w = 8.5 - left_m - right_m

    # -- Set default paragraph style --
    style = doc.styles['Normal']
    style.font.name = _font_name
    style.font.size = Pt(_body_size)
    style.font.color.rgb = COLOR_DARK
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)
    if s["line_spacing"]:
        style.paragraph_format.line_spacing = s["line_spacing"]

    # ==========================================
    # HEADER — Name and contact info
    # ==========================================
    _add_letterhead_docx(doc, profile, s, content_w, include_links=True)

    # ==========================================
    # PARSE AI CONTENT
    # ==========================================
    sections = _parse_resume_sections(content)

    def _section_header(title: str):
        """Add a section header respecting the chosen style."""
        para = doc.add_paragraph()
        _set_paragraph_spacing(para, before=s["section_gap"], after=3)
        # Small caps keeps the mixed-case text and renders it as caps; the
        # other styles upper-case the text outright.
        run = para.add_run(title if s["header_small_caps"] else title.upper())
        _set_font(run, name=_font_name, size=_header_size, bold=True,
                  color=_rgb(s["header_color"]))
        if s["header_small_caps"]:
            run.font.small_caps = True
        else:
            run.font.all_caps = True
        if s["header_letter_spacing"]:
            run.font.letter_spacing = Pt(s["header_letter_spacing"])
        # Never leave a header stranded at the foot of a page.
        para.paragraph_format.keep_with_next = True
        if s["header_underline"]:
            if s["header_rule_style"] == "stub":
                stub = _add_stub_rule(doc, content_w, 0.30,
                                      s["header_underline_color"],
                                      s["header_underline_size"],
                                      before=1, after=4)
                stub.paragraph_format.keep_with_next = True
            else:
                _add_bottom_border(para, color=s["header_underline_color"],
                                   size=s["header_underline_size"])
        return para

    # ==========================================
    # SUMMARY
    # ==========================================
    summary_text = sections.get("summary", "")
    if summary_text:
        _section_header("Summary")
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=2, after=4)
        # Clean up any stray markdown
        clean_summary = re.sub(r'\*\*(.+?)\*\*', r'\1', summary_text).strip()
        run = p.add_run(clean_summary)
        _set_font(run, name=_font_name, size=_body_size, color=COLOR_DARK)

    # ==========================================
    # SKILLS
    # ==========================================
    skills_text = sections.get("skills", "")
    if skills_text:
        _section_header("Technical Skills")

        # Parse skills — could be comma-separated, bullet list, or categorized
        skills_text = re.sub(r'\*\*(.+?)\*\*', r'\1', skills_text)
        # Remove bullet prefixes
        skills_text = re.sub(r'^[-*•]\s*', '', skills_text, flags=re.MULTILINE)
        # Collect non-empty lines
        skill_lines = [line.strip() for line in skills_text.split("\n") if line.strip()]

        # Check if skills are categorized (e.g., "Security: skill1, skill2")
        categorized = any(":" in line for line in skill_lines)

        if categorized:
            for line in skill_lines:
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, before=1, after=1)
                if ":" in line:
                    category, skills = line.split(":", 1)
                    cat_run = p.add_run(category.strip() + ": ")
                    _set_font(cat_run, name=_font_name, size=_body_size, bold=True, color=COLOR_DARK)
                    skills_run = p.add_run(skills.strip())
                    _set_font(skills_run, name=_font_name, size=_body_size, color=COLOR_DARK)
                else:
                    run = p.add_run(line)
                    _set_font(run, name=_font_name, size=_body_size, color=COLOR_DARK)
        else:
            # Join all onto one line with the style's separator
            all_skills: list[str] = []
            for line in skill_lines:
                all_skills.extend(sk.strip() for sk in line.split(",") if sk.strip())
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=2, after=4)
            run = p.add_run(s["skills_separator"].join(all_skills))
            _set_font(run, name=_font_name, size=_body_size, color=COLOR_DARK)

    # ==========================================
    # EXPERIENCE
    # ==========================================
    exp_text = sections.get("experience", "")
    if exp_text:
        _section_header("Professional Experience")
        entries = _parse_experience_entries(exp_text)

        # Inline layout uses full text-width tab; stacked keeps the legacy 6.1" stop.
        inline_right_tab = Inches(content_w)
        stacked_right_tab = Inches(6.1)
        title_size = _body_size + 0.5
        company_accent = s["company_style"] == "accent"
        marker_color = _rgb(s["bullet_marker_color"])

        def _add_company_run(para):
            """Title/company separator + company name, per the style's
            company_style token."""
            sep_run = para.add_run(s["entry_separator"])
            _set_font(sep_run, name=_font_name, size=_body_size,
                      color=_rgb(_accent_hex) if company_accent else COLOR_GRAY)
            company_run = para.add_run(entry["company"])
            _set_font(
                company_run, name=_font_name, size=_body_size,
                bold=company_accent,
                italic=s["company_style"] == "italic",
                color=_rgb(_accent_hex) if company_accent else COLOR_DARK,
            )

        for entry in entries:
            if s["entry_layout"] == "inline":
                # One-line header: "Title  ·  Company                Dates"
                head = doc.add_paragraph()
                _set_paragraph_spacing(head, before=6, after=2)
                head.paragraph_format.keep_with_next = True

                title_run = head.add_run(entry["title"])
                _set_font(title_run, name=_font_name, size=title_size, bold=True, color=COLOR_BLACK)

                if entry["company"]:
                    _add_company_run(head)

                if entry["dates"]:
                    head.paragraph_format.tab_stops.add_tab_stop(
                        inline_right_tab, alignment=WD_ALIGN_PARAGRAPH.RIGHT
                    )
                    head.add_run("\t")
                    dates_run = head.add_run(entry["dates"])
                    _set_font(dates_run, name=_font_name, size=_dates_size, italic=True, color=COLOR_GRAY)
            else:
                # Stacked: title/dates on line 1, company in italic on line 2
                title_para = doc.add_paragraph()
                _set_paragraph_spacing(title_para, before=6, after=0)
                title_para.paragraph_format.keep_with_next = True

                title_run = title_para.add_run(entry["title"])
                _set_font(title_run, name=_font_name, size=title_size, bold=True, color=COLOR_BLACK)

                if entry["dates"]:
                    title_para.paragraph_format.tab_stops.add_tab_stop(
                        stacked_right_tab, alignment=WD_ALIGN_PARAGRAPH.RIGHT
                    )
                    title_para.add_run("\t")
                    dates_run = title_para.add_run(entry["dates"])
                    _set_font(dates_run, name=_font_name, size=_dates_size, italic=True, color=COLOR_GRAY)

                if entry["company"]:
                    company_para = doc.add_paragraph()
                    _set_paragraph_spacing(company_para, before=0, after=2)
                    company_para.paragraph_format.keep_with_next = True
                    company_run = company_para.add_run(entry["company"])
                    _set_font(company_run, name=_font_name, size=_body_size, italic=True, color=COLOR_GRAY)

            # Bullets — marker glyph and color come from the style preset
            for bullet in entry["bullets"]:
                bullet_para = doc.add_paragraph()
                _set_paragraph_spacing(bullet_para, before=1, after=1)
                bullet_para.paragraph_format.left_indent = Inches(0.25)
                bullet_para.paragraph_format.first_line_indent = Inches(-0.2)

                marker_run = bullet_para.add_run(s["bullet_marker"])
                _set_font(marker_run, name=_font_name, size=s["bullet_marker_size"], color=marker_color)
                text_run = bullet_para.add_run(bullet)
                _set_font(text_run, name=_font_name, size=_body_size, color=COLOR_DARK)

    # ==========================================
    # EDUCATION
    # ==========================================
    edu_text = sections.get("education", "")
    if edu_text:
        _section_header("Education")
        entries = _parse_education_entries(edu_text)

        for entry in entries:
            edu_para = doc.add_paragraph()
            _set_paragraph_spacing(edu_para, before=3, after=1)

            degree_run = edu_para.add_run(entry["degree"])
            _set_font(degree_run, name=_font_name, size=_body_size, bold=True, color=COLOR_BLACK)

            if entry["school"]:
                school_text = f"  —  {entry['school']}"
                if entry["year"]:
                    school_text += f"  ({entry['year']})"
                school_run = edu_para.add_run(school_text)
                _set_font(school_run, name=_font_name, size=_body_size, color=COLOR_GRAY)
            elif entry["year"]:
                year_run = edu_para.add_run(f"  ({entry['year']})")
                _set_font(year_run, name=_font_name, size=_body_size, color=COLOR_GRAY)

    # ==========================================
    # CERTIFICATIONS
    # ==========================================
    cert_text = sections.get("certifications", "")
    if cert_text:
        _section_header("Certifications")
        for line in cert_text.split("\n"):
            line = line.strip()
            if not line:
                continue
            line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
            clean = line.lstrip("-*•– ").strip()
            if clean:
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, before=1, after=1)
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.first_line_indent = Inches(-0.2)
                marker_run = p.add_run(s["bullet_marker"])
                _set_font(marker_run, name=_font_name, size=s["bullet_marker_size"], color=_rgb(s["bullet_marker_color"]))
                run = p.add_run(clean)
                _set_font(run, name=_font_name, size=_body_size, color=COLOR_DARK)

    # ==========================================
    # REFERENCES — appended verbatim from profile, never sent to AI
    # ==========================================
    references = profile.get("references", []) or []
    valid_refs = [r for r in references if (r.get("name") or "").strip()]
    if valid_refs:
        _section_header("References")
        for ref in valid_refs:
            name_para = doc.add_paragraph()
            _set_paragraph_spacing(name_para, before=4, after=1)
            name_run = name_para.add_run(ref.get("name", "").strip())
            _set_font(name_run, name=_font_name, size=_body_size, bold=True, color=COLOR_BLACK)
            position = (ref.get("position") or "").strip()
            if position:
                pos_run = name_para.add_run(f"  —  {position}")
                _set_font(pos_run, name=_font_name, size=_body_size, italic=True, color=COLOR_GRAY)

            contact_bits = []
            email = (ref.get("email") or "").strip()
            phone = (ref.get("phone") or "").strip()
            if email:
                contact_bits.append(email)
            if phone:
                contact_bits.append(phone)
            if contact_bits:
                contact_para = doc.add_paragraph()
                _set_paragraph_spacing(contact_para, before=0, after=2)
                run = contact_para.add_run(s["contact_separator"].join(contact_bits))
                _set_font(run, name=_font_name, size=_dates_size, color=COLOR_DARK)

    # -- Save --
    job_id = job.get("id", "unknown")
    file_path = RESUMES_DIR / f"{job_id}_resume.docx"
    doc.save(str(file_path))
    logger.info("Resume saved to %s", file_path)
    return str(file_path)


def generate_cover_letter_docx(content: str, profile: dict, job: dict, config: dict | None = None) -> str:
    """
    Create a professional cover letter DOCX in the same visual style as the
    resume, so a recruiter opening both files sees one matched set.
    Returns the file path to the generated document.
    """
    _ensure_dir()
    s = _resolve_resume_style(config)
    body_font = s["body_font"]
    doc = Document()

    # A letter always breathes more than a resume — 1" margins regardless of
    # the preset's tighter resume margins.
    left_m = right_m = 1.0
    for section in doc.sections:
        section.top_margin = Inches(0.8 if s["banner"] else 1.0)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(left_m)
        section.right_margin = Inches(right_m)

    content_w = 8.5 - left_m - right_m

    # Default style
    style = doc.styles['Normal']
    style.font.name = body_font
    style.font.size = Pt(11)
    style.font.color.rgb = COLOR_DARK
    if s["line_spacing"]:
        style.paragraph_format.line_spacing = s["line_spacing"]

    # Candidate letterhead — identical treatment to the resume
    _add_letterhead_docx(doc, profile, s, content_w, include_links=False)

    # Date
    date_para = doc.add_paragraph()
    _set_paragraph_spacing(date_para, before=10, after=12)
    date_run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
    _set_font(date_run, name=body_font, size=11, color=COLOR_GRAY)

    # Role reference
    ref_para = doc.add_paragraph()
    _set_paragraph_spacing(ref_para, before=0, after=12)
    ref_run = ref_para.add_run(f"Re: {job.get('title', '')} at {job.get('company', '')}")
    _set_font(ref_run, name=body_font, size=11, bold=True, color=COLOR_BLACK)

    # Body paragraphs
    # Clean up any markdown the AI might have included
    clean_content = re.sub(r'\*\*(.+?)\*\*', r'\1', content)
    clean_content = re.sub(r'^[#]+\s*', '', clean_content, flags=re.MULTILINE)

    for paragraph in clean_content.split("\n\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        # Skip closings the AI added — we add our own
        if paragraph.lower().startswith(("sincerely", "best regards", "regards", "thank you")):
            continue
        p = doc.add_paragraph()
        _set_paragraph_spacing(p, before=0, after=6)
        run = p.add_run(paragraph)
        _set_font(run, name=body_font, size=11, color=COLOR_DARK)

    # Closing
    doc.add_paragraph()
    closing = doc.add_paragraph()
    _set_paragraph_spacing(closing, before=6, after=2)
    run = closing.add_run("Sincerely,")
    _set_font(run, name=body_font, size=11, color=COLOR_DARK)

    name_para = doc.add_paragraph()
    _set_paragraph_spacing(name_para, before=2, after=0)
    run = name_para.add_run(profile.get("full_name", ""))
    _set_font(run, name=body_font, size=11, bold=True, color=COLOR_BLACK)

    # Save
    job_id = job.get("id", "unknown")
    file_path = RESUMES_DIR / f"{job_id}_cover_letter.docx"
    doc.save(str(file_path))
    logger.info("Cover letter saved to %s", file_path)
    return str(file_path)


# ===========================================================================
# PDF rendering (pure-Python via ReportLab — bundled, no external binary)
#
# The PDF path reuses the exact same parsers (_parse_resume_sections,
# _parse_experience_entries, _parse_education_entries) and the _STYLES presets
# the DOCX path uses, so the two formats stay visually consistent.
# ===========================================================================


def _document_format(config: dict | None) -> str:
    """Return the configured output format: 'pdf' or 'docx' (default)."""
    if not config:
        return "docx"
    fmt = config.get("application_honesty", {}).get("document_format", "docx")
    return "pdf" if str(fmt).lower() == "pdf" else "docx"


def _pdf_esc(text: str) -> str:
    """Escape text for ReportLab's mini-markup Paragraph parser."""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


_PDF_SAFE_BULLET = "•"
# Glyphs the built-in WinAnsi fonts (and our embedded faces) can all render.
_PDF_SAFE_MARKS = {"•", "–", "—", "·"}


def _pdf_bullet(marker: str) -> str:
    """Decorative glyphs like ▸/▪/‣ render as a .notdef box in the built-in
    PDF fonts. Keep ASCII and the known-safe marks; map anything else to a
    plain round bullet."""
    return "".join(
        ch if (ch.isascii() or ch in _PDF_SAFE_MARKS) else _PDF_SAFE_BULLET
        for ch in marker
    )


# -- Embedded PDF typefaces (OFL) -------------------------------------------
# ReportLab's built-in fonts are Helvetica/Times only, which is why PDFs used
# to look dated next to the DOCX. We ship Lato (sans) and PT Serif (serif) and
# register them on first use. Embedded TTFs keep the text fully extractable —
# ATS parsing is unaffected. If the files are missing (or ReportLab rejects
# them) every style silently falls back to the built-ins.
FONTS_DIR = Path(__file__).resolve().parent / "fonts"

_EMBEDDED_FAMILIES: dict[str, dict[str, str]] = {
    "JobsmithSans":  {
        "normal": "Lato-Regular", "bold": "Lato-Bold",
        "italic": "Lato-Italic", "boldItalic": "Lato-BoldItalic",
    },
    "JobsmithSerif": {
        "normal": "PTSerif-Regular", "bold": "PTSerif-Bold",
        "italic": "PTSerif-Italic", "boldItalic": "PTSerif-BoldItalic",
    },
}

_embedded_fonts_ready: bool | None = None  # None = not attempted yet


def _register_embedded_fonts() -> bool:
    """Register the bundled TTFs with ReportLab once. Returns True when the
    embedded families are available."""
    global _embedded_fonts_ready
    if _embedded_fonts_ready is not None:
        return _embedded_fonts_ready

    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.lib.fonts import addMapping

        for family, faces in _EMBEDDED_FAMILIES.items():
            for variant, stem in faces.items():
                path = FONTS_DIR / f"{stem}.ttf"
                if not path.exists():
                    raise FileNotFoundError(path)
                face = family if variant == "normal" else f"{family}-{variant}"
                pdfmetrics.registerFont(TTFont(face, str(path)))
            # Teach ReportLab which face to use for <b>/<i> markup.
            pdfmetrics.registerFontFamily(
                family,
                normal=family,
                bold=f"{family}-bold",
                italic=f"{family}-italic",
                boldItalic=f"{family}-boldItalic",
            )
            for bold, italic, variant in (
                (0, 0, "normal"), (1, 0, "bold"),
                (0, 1, "italic"), (1, 1, "boldItalic"),
            ):
                face = family if variant == "normal" else f"{family}-{variant}"
                addMapping(family, bold, italic, face)

        _embedded_fonts_ready = True
        logger.info("Embedded PDF fonts registered (Lato, PT Serif)")
    except Exception:
        _embedded_fonts_ready = False
        logger.warning(
            "Embedded PDF fonts unavailable — falling back to built-in "
            "Helvetica/Times", exc_info=True,
        )
    return _embedded_fonts_ready


def _pdf_font(style_font: str, bold: bool = False, italic: bool = False) -> str:
    """Map a preset font name to a concrete PDF face.

    Swiss deliberately renders in Helvetica — the built-in grotesque *is* the
    design — so its Arial token maps straight to the built-in. Everything else
    prefers the embedded families and degrades to the built-ins.
    """
    font = style_font.lower()

    if "arial" in font or "helvetica" in font:
        family = None  # force built-in
    elif "georgia" in font or "times" in font or "serif" in font:
        family = "JobsmithSerif"
    else:
        family = "JobsmithSans"

    if family and _register_embedded_fonts():
        if bold and italic:
            return f"{family}-boldItalic"
        if bold:
            return f"{family}-bold"
        if italic:
            return f"{family}-italic"
        return family

    # Built-in fallback
    if family == "JobsmithSerif" or "times" in font or "georgia" in font:
        if bold and italic:
            return "Times-BoldItalic"
        if bold:
            return "Times-Bold"
        if italic:
            return "Times-Italic"
        return "Times-Roman"
    if bold and italic:
        return "Helvetica-BoldOblique"
    if bold:
        return "Helvetica-Bold"
    if italic:
        return "Helvetica-Oblique"
    return "Helvetica"


def _rl_color(rgb):
    """Convert a python-docx RGBColor (3 ints) to a ReportLab color."""
    from reportlab.lib.colors import HexColor
    return HexColor("#%02X%02X%02X" % tuple(rgb))


def _rl_hex(hex6: str):
    """Convert a bare 6-char hex string ('2B5797') to a ReportLab color."""
    from reportlab.lib.colors import HexColor
    return HexColor("#" + hex6)


def _border_pt(eighths: str) -> float:
    """DOCX border sizes are in 1/8 pt units; convert to points for ReportLab."""
    try:
        return max(0.4, float(eighths) / 8.0)
    except (TypeError, ValueError):
        return 0.75


def _pdf_para(text, font, size, color, *, align=None, bold=False, italic=False,
              sb=0, sa=2, leading=None, left_indent=0, first_indent=0,
              keep=False, line_spacing=None):
    """One Paragraph in the resolved style's typeface."""
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import Paragraph

    st = ParagraphStyle(
        "x", fontName=_pdf_font(font, bold, italic), fontSize=size,
        leading=leading or size * 1.18 * (line_spacing or 1.0),
        textColor=color, alignment=TA_LEFT if align is None else align,
        spaceBefore=sb, spaceAfter=sa, leftIndent=left_indent,
        firstLineIndent=first_indent, keepWithNext=1 if keep else 0,
    )
    return Paragraph(text, st)


def _letterhead_story(profile: dict, s: dict, avail_w: float, *,
                      include_links: bool) -> list:
    """Name + contact letterhead flowables — shared by the resume and cover
    letter PDF paths so a generated pair looks like one set."""
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib.units import inch
    from reportlab.platypus import Table, TableStyle, HRFlowable, Spacer

    align = TA_CENTER if s["name_align"] == "center" else TA_LEFT
    body = s["body_font"]
    body_size = s["body_size"]
    contact_size = max(8.0, body_size - 1.5)
    banner = s["banner"]
    story: list = []

    name = profile.get("full_name", "")
    name_flow = None
    if name:
        # Small caps has no PDF equivalent in the built-in markup — upper-case
        # with the preset's letter-spacing reads the same at a glance.
        disp = name.upper() if (s["name_uppercase"] or s["name_small_caps"]) else name
        name_flow = _pdf_para(
            _pdf_esc(disp), s["name_font"], s["name_size"],
            _rl_hex(s["name_color"]), align=align, bold=s["name_bold"], sa=2,
        )

    contact_bits: list[str] = []
    link_hex = "#" + s["accent_color"]
    for text, url in _contact_entries(profile, include_links=include_links):
        if url and s["hyperlinks"] and not banner:
            contact_bits.append(
                f'<a href="{_pdf_esc(url)}" color="{link_hex}">{_pdf_esc(text)}</a>'
            )
        else:
            contact_bits.append(_pdf_esc(text))

    contact_flow = None
    if contact_bits:
        contact_flow = _pdf_para(
            s["contact_separator"].join(contact_bits), body, contact_size,
            _rl_hex(s["contact_color"]), align=align, sa=2,
        )

    # Compact puts the contact line on the name line itself.
    if s["contact_inline"] and name_flow is not None and contact_bits:
        merged = (
            f'<font size="{s["name_size"]}" color="#{s["name_color"]}">'
            f'<b>{_pdf_esc(name)}</b></font>'
            f'<font size="{contact_size}" color="#{s["contact_color"]}">'
            f'   {s["contact_separator"].join(contact_bits)}</font>'
        )
        story.append(_pdf_para(merged, s["name_font"], s["name_size"],
                               _rl_hex(s["name_color"]), align=align, sa=2))
        name_flow = contact_flow = None

    if banner:
        # The band: a single-cell table with a solid background. Text inside is
        # ordinary text — extraction and ATS parsing are unaffected.
        rows = [[f] for f in (name_flow, contact_flow) if f is not None]
        if rows:
            t = Table(rows, colWidths=[avail_w])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), _rl_hex(s["accent"])),
                ("LEFTPADDING", (0, 0), (-1, -1), 12),
                ("RIGHTPADDING", (0, 0), (-1, -1), 12),
                ("TOPPADDING", (0, 0), (0, 0), 12),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, -1), (-1, -1), 12),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(t)
            story.append(Spacer(1, 10))
        return story

    for flow in (name_flow, contact_flow):
        if flow is not None:
            story.append(flow)

    if s["name_rule"]:
        rule_style = s["name_rule_style"]
        thickness = _border_pt(s["name_rule_size"])
        color = _rl_hex(s["name_rule_color"])
        if rule_style == "stub":
            story.append(HRFlowable(
                width=0.55 * inch, thickness=thickness, color=color,
                hAlign="LEFT", spaceBefore=4, spaceAfter=7,
            ))
        elif rule_style == "double":
            story.append(HRFlowable(width="100%", thickness=thickness,
                                    color=color, spaceBefore=5, spaceAfter=1.5))
            story.append(HRFlowable(width="100%", thickness=thickness,
                                    color=color, spaceBefore=0, spaceAfter=5))
        else:
            story.append(HRFlowable(width="100%", thickness=thickness,
                                    color=color, spaceBefore=4, spaceAfter=4))
    return story


def _render_resume_pdf(content: str, profile: dict, job: dict, config: dict | None = None) -> str:
    """Render the resume to a styled, ATS-friendly PDF mirroring the DOCX
    layout. Returns the PDF file path."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_RIGHT
    from reportlab.platypus import (
        SimpleDocTemplate, Table, TableStyle, HRFlowable,
    )

    _ensure_dir()
    s = _resolve_resume_style(config)
    bullet_marker = _pdf_bullet(s["bullet_marker"])
    body = s["body_font"]
    body_size = s["body_size"]
    dates_size = max(8.0, body_size - 0.5)
    ls = s["line_spacing"]
    top_m, bot_m, left_m, right_m = s["margins"]

    job_id = job.get("id", "unknown")
    file_path = RESUMES_DIR / f"{job_id}_resume.pdf"

    doc = SimpleDocTemplate(
        str(file_path), pagesize=LETTER,
        topMargin=top_m * inch, bottomMargin=bot_m * inch,
        leftMargin=left_m * inch, rightMargin=right_m * inch,
        title=f"{profile.get('full_name', 'Resume')} — Resume",
    )
    avail_w = LETTER[0] - (left_m + right_m) * inch

    accent_hex = "#" + s["accent_color"]
    c_dark = _rl_color(COLOR_DARK)
    c_gray = _rl_color(COLOR_GRAY)
    c_black = _rl_color(COLOR_BLACK)
    story: list = []

    def para(text, font, size, color, **kw):
        kw.setdefault("line_spacing", ls)
        return _pdf_para(text, font, size, color, **kw)

    # -- Letterhead (name, contact, rule/band) --
    story.extend(_letterhead_story(profile, s, avail_w, include_links=True))

    sections = _parse_resume_sections(content)

    def section_header(title: str):
        story.append(para(
            _pdf_esc(title.upper()), body, s["header_size"],
            _rl_hex(s["header_color"]), bold=True,
            sb=s["section_gap"], sa=3, keep=True,
        ))
        if s["header_underline"]:
            width = (0.30 * inch if s["header_rule_style"] == "stub" else "100%")
            story.append(HRFlowable(
                width=width, thickness=_border_pt(s["header_underline_size"]),
                color=_rl_hex(s["header_underline_color"]), hAlign="LEFT",
                spaceBefore=1, spaceAfter=4,
            ))

    # -- Summary --
    summary_text = sections.get("summary", "")
    if summary_text:
        section_header("Summary")
        clean = re.sub(r"\*\*(.+?)\*\*", r"\1", summary_text).strip()
        story.append(para(_pdf_esc(clean), body, body_size, c_dark, sb=2, sa=4))

    # -- Skills --
    skills_text = sections.get("skills", "")
    if skills_text:
        section_header("Technical Skills")
        skills_text = re.sub(r"\*\*(.+?)\*\*", r"\1", skills_text)
        skills_text = re.sub(r"^[-*•]\s*", "", skills_text, flags=re.MULTILINE)
        skill_lines = [ln.strip() for ln in skills_text.split("\n") if ln.strip()]
        if any(":" in ln for ln in skill_lines):
            for ln in skill_lines:
                if ":" in ln:
                    cat, sk = ln.split(":", 1)
                    txt = f"<b>{_pdf_esc(cat.strip())}:</b> {_pdf_esc(sk.strip())}"
                else:
                    txt = _pdf_esc(ln)
                story.append(para(txt, body, body_size, c_dark, sb=1, sa=1))
        else:
            all_skills: list[str] = []
            for ln in skill_lines:
                all_skills.extend(p.strip() for p in ln.split(",") if p.strip())
            joined = _pdf_esc(s["skills_separator"]).join(
                _pdf_esc(sk) for sk in all_skills
            )
            story.append(para(joined, body, body_size, c_dark, sb=2, sa=4))

    def _no_border_table(left_flow, right_flow, right_w):
        t = Table([[left_flow, right_flow]],
                  colWidths=[avail_w - right_w, right_w])
        t.setStyle(TableStyle([
            ("LEFTPADDING", (0, 0), (-1, -1), 0),
            ("RIGHTPADDING", (0, 0), (-1, -1), 0),
            ("TOPPADDING", (0, 0), (-1, -1), 0),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        return t

    # -- Experience --
    exp_text = sections.get("experience", "")
    if exp_text:
        section_header("Professional Experience")
        right_w = 1.7 * inch
        title_size = body_size + 0.5
        company_accent = s["company_style"] == "accent"
        sep = _pdf_esc(s["entry_separator"])
        for entry in _parse_experience_entries(exp_text):
            if s["entry_layout"] == "inline":
                left = f'<b>{_pdf_esc(entry["title"])}</b>'
                if entry["company"]:
                    company = _pdf_esc(entry["company"])
                    if company_accent:
                        left += (f'<font color="{accent_hex}">{sep}'
                                 f'<b>{company}</b></font>')
                    elif s["company_style"] == "italic":
                        left += f'<font color="#666666">{sep}</font><i>{company}</i>'
                    else:
                        left += f'<font color="#666666">{sep}</font>{company}'
                left_p = para(left, body, title_size, c_black, sb=6, sa=2, keep=True)
                if entry["dates"]:
                    right_p = para(_pdf_esc(entry["dates"]), body, dates_size,
                                   c_gray, align=TA_RIGHT, italic=True, sb=6, sa=2)
                    story.append(_no_border_table(left_p, right_p, right_w))
                else:
                    story.append(left_p)
            else:
                title_p = para(f'<b>{_pdf_esc(entry["title"])}</b>', body,
                               title_size, c_black, sb=6, sa=0, keep=True)
                if entry["dates"]:
                    right_p = para(_pdf_esc(entry["dates"]), body, dates_size,
                                   c_gray, align=TA_RIGHT, italic=True, sb=6, sa=0)
                    story.append(_no_border_table(title_p, right_p, right_w))
                else:
                    story.append(title_p)
                if entry["company"]:
                    story.append(para(_pdf_esc(entry["company"]), body, body_size,
                                      c_gray, italic=True, sb=0, sa=2, keep=True))
            marker_hex = "#" + s["bullet_marker_color"]
            for bullet in entry["bullets"]:
                txt = (f'<font color="{marker_hex}">'
                       f'{_pdf_esc(bullet_marker)}</font>'
                       f'{_pdf_esc(bullet)}')
                story.append(para(txt, body, body_size, c_dark, sb=1, sa=1,
                                  left_indent=18, first_indent=-12))

    # -- Education --
    edu_text = sections.get("education", "")
    if edu_text:
        section_header("Education")
        for entry in _parse_education_entries(edu_text):
            txt = f'<b>{_pdf_esc(entry["degree"])}</b>'
            if entry["school"]:
                tail = f'  —  {entry["school"]}'
                if entry["year"]:
                    tail += f'  ({entry["year"]})'
                txt += f'<font color="#666666">{_pdf_esc(tail)}</font>'
            elif entry["year"]:
                txt += f'<font color="#666666">  ({_pdf_esc(entry["year"])})</font>'
            story.append(para(txt, body, body_size, c_black, sb=3, sa=1))

    # -- Certifications --
    cert_text = sections.get("certifications", "")
    if cert_text:
        section_header("Certifications")
        marker_hex = "#" + s["bullet_marker_color"]
        for line in cert_text.split("\n"):
            line = re.sub(r"\*\*(.+?)\*\*", r"\1", line.strip())
            clean = line.lstrip("-*•– ").strip()
            if clean:
                txt = (f'<font color="{marker_hex}">'
                       f'{_pdf_esc(bullet_marker)}</font>{_pdf_esc(clean)}')
                story.append(para(txt, body, body_size, c_dark, sb=1, sa=1,
                                  left_indent=18, first_indent=-12))

    # -- References (verbatim from profile, never AI-touched) --
    refs = [r for r in (profile.get("references", []) or [])
            if (r.get("name") or "").strip()]
    if refs:
        section_header("References")
        for ref in refs:
            line = f'<b>{_pdf_esc(ref["name"].strip())}</b>'
            position = (ref.get("position") or "").strip()
            if position:
                line += f'<i><font color="#666666">  —  {_pdf_esc(position)}</font></i>'
            story.append(para(line, body, body_size, c_black, sb=4, sa=1))
            bits = [b for b in ((ref.get("email") or "").strip(),
                                (ref.get("phone") or "").strip()) if b]
            if bits:
                story.append(para(
                    _pdf_esc(s["contact_separator"].join(bits)), body,
                    dates_size, c_dark, sb=0, sa=2,
                ))

    doc.build(story)
    logger.info("Resume PDF saved to %s", file_path)
    return str(file_path)


def _render_cover_letter_pdf(content: str, profile: dict, job: dict,
                             config: dict | None = None) -> str:
    """Render the cover letter to PDF in the resume's visual style, mirroring
    generate_cover_letter_docx."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Spacer

    _ensure_dir()
    s = _resolve_resume_style(config)
    body = s["body_font"]
    ls = s["line_spacing"]

    job_id = job.get("id", "unknown")
    file_path = RESUMES_DIR / f"{job_id}_cover_letter.pdf"

    top_m = 0.8 if s["banner"] else 1.0
    doc = SimpleDocTemplate(
        str(file_path), pagesize=LETTER,
        topMargin=top_m * inch, bottomMargin=inch,
        leftMargin=inch, rightMargin=inch,
        title=f"{profile.get('full_name', 'Cover Letter')} — Cover Letter",
    )
    avail_w = LETTER[0] - 2 * inch

    c_dark = _rl_color(COLOR_DARK)
    c_gray = _rl_color(COLOR_GRAY)
    c_black = _rl_color(COLOR_BLACK)
    story: list = []

    def para(text, size, color, **kw):
        kw.setdefault("line_spacing", ls)
        kw.setdefault("sa", 6)
        return _pdf_para(text, body, size, color, **kw)

    story.extend(_letterhead_story(profile, s, avail_w, include_links=False))

    story.append(para(_pdf_esc(datetime.now().strftime("%B %d, %Y")), 11,
                      c_gray, sb=10, sa=12))
    story.append(para(
        f'<b>Re: {_pdf_esc(job.get("title", ""))} at '
        f'{_pdf_esc(job.get("company", ""))}</b>', 11, c_black, sa=12))

    clean = re.sub(r"\*\*(.+?)\*\*", r"\1", content)
    clean = re.sub(r"^[#]+\s*", "", clean, flags=re.MULTILINE)
    for paragraph in clean.split("\n\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        low = paragraph.lower()
        if low.startswith(("sincerely", "best regards", "regards", "thank you")):
            continue  # we add our own closing
        story.append(para(_pdf_esc(paragraph), 11, c_dark, sa=6))

    story.append(Spacer(1, 12))
    story.append(para("Sincerely,", 11, c_dark, sb=6, sa=2))
    story.append(para(f'<b>{_pdf_esc(profile.get("full_name", ""))}</b>', 11,
                      c_black, sb=2, sa=0))

    doc.build(story)
    logger.info("Cover letter PDF saved to %s", file_path)
    return str(file_path)


# ---------------------------------------------------------------------------
# Public dispatchers — pick format from config.application_honesty.document_format
# ---------------------------------------------------------------------------


def generate_resume(content: str, profile: dict, job: dict, config: dict | None = None) -> str:
    """Generate the resume in the user's configured format.

    Always writes the DOCX (the reliable internal artifact the autofill
    adapters and Applicant Assist drag/drop flow rely on). When the format is
    "pdf", also renders a PDF and returns its path; if PDF rendering fails,
    falls back to the DOCX so the pipeline never breaks.
    """
    docx_path = generate_resume_docx(content, profile, job, config)
    if _document_format(config) == "pdf":
        try:
            return _render_resume_pdf(content, profile, job, config)
        except Exception:
            logger.warning(
                "Resume PDF rendering failed — serving DOCX instead", exc_info=True
            )
    return docx_path


def generate_cover_letter(content: str, profile: dict, job: dict, config: dict | None = None) -> str:
    """Generate the cover letter in the user's configured format.

    Always writes the DOCX; additionally renders a PDF when configured,
    falling back to the DOCX if PDF rendering fails.
    """
    docx_path = generate_cover_letter_docx(content, profile, job, config)
    if _document_format(config) == "pdf":
        try:
            return _render_cover_letter_pdf(content, profile, job, config)
        except Exception:
            logger.warning(
                "Cover letter PDF rendering failed — serving DOCX instead",
                exc_info=True,
            )
    return docx_path
