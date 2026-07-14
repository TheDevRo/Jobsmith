"""
Tests for backend/resume_generator.py style presets.

Covers the five-style lineup (executive / ledger / banner / compact / swiss),
the accent-color setting, legacy style-name migration, and — most importantly —
the ATS invariant: every style must stay single-column real text, with no
tables, images or text boxes in the DOCX.
"""

import zipfile

import pytest
from docx import Document

from backend import resume_generator as rg


STYLES = ["executive", "ledger", "banner", "compact", "swiss"]

PROFILE = {
    "full_name": "Morgan Reyes",
    "email": "morgan.reyes@example.com",
    "phone": "(503) 555-0142",
    "location": "Portland, OR",
    "linkedin": "linkedin.com/in/morganreyes",
    "references": [
        {"name": "Dana Whitfield", "position": "VP Analytics",
         "email": "dana.w@example.com", "phone": "(503) 555-0188"},
    ],
}

RESUME_CONTENT = """SUMMARY
Data analyst who turns messy operational data into decisions leaders act on.

TECHNICAL SKILLS
SQL, Python, dbt, Snowflake, Tableau

PROFESSIONAL EXPERIENCE
Title: Senior Data Analyst
Company: Northwind Logistics
Dates: 2021 - Present
- Cut freight spend 12% ($2.1M annually) across 14 distribution centers.

EDUCATION
Degree: B.S. Statistics
School: Oregon State University
Year: 2016

CERTIFICATIONS
- Tableau Desktop Certified Professional
"""

LETTER_CONTENT = "Dear Hiring Team,\n\nI would like to apply.\n\nSincerely,\nMorgan Reyes\n"

JOB = {"id": "t1", "title": "Senior Analytics Engineer", "company": "Meridian Freight"}


def _cfg(style="ledger", accent="default", fmt="docx"):
    return {"application_honesty": {
        "resume_style": style, "resume_accent": accent, "document_format": fmt,
    }}


def _document_xml(path) -> str:
    with zipfile.ZipFile(path) as z:
        return z.read("word/document.xml").decode("utf-8")


# ---------------------------------------------------------------------------
# Style resolution
# ---------------------------------------------------------------------------
@pytest.mark.parametrize("style", STYLES)
def test_every_style_resolves_with_full_token_set(style):
    s = rg._resolve_resume_style(_cfg(style))
    # Presets only carry overrides — resolution must fill in every default.
    for token in rg._STYLE_DEFAULTS:
        assert token in s, f"{style} missing token {token}"


@pytest.mark.parametrize("legacy,expected", [
    ("standard", "ledger"),
    ("modern", "ledger"),
    ("minimal", "swiss"),
])
def test_legacy_style_names_still_resolve(legacy, expected):
    """Configs written before the five-style lineup must keep working."""
    got = rg._resolve_resume_style(_cfg(legacy))
    assert got == rg._resolve_resume_style(_cfg(expected))


def test_unknown_style_falls_back_to_default():
    got = rg._resolve_resume_style(_cfg("nonsense"))
    assert got == rg._resolve_resume_style(_cfg(rg.DEFAULT_RESUME_STYLE))


def test_missing_config_uses_default_style():
    assert rg._resolve_resume_style(None) == rg._resolve_resume_style(_cfg("ledger"))


# ---------------------------------------------------------------------------
# Accent colors
# ---------------------------------------------------------------------------
def test_accent_choice_recolors_accent_driven_styles():
    s = rg._resolve_resume_style(_cfg("ledger", accent="burgundy"))
    assert s["accent"] == rg.ACCENT_CHOICES["burgundy"]
    # Tokens carrying the "accent" sentinel resolve to the chosen hex.
    assert s["header_color"] == rg.ACCENT_CHOICES["burgundy"]
    assert s["name_rule_color"] == rg.ACCENT_CHOICES["burgundy"]


@pytest.mark.parametrize("style", ["executive", "swiss"])
def test_monochrome_styles_ignore_accent_choice(style):
    """Executive and Swiss are deliberately monochrome."""
    default = rg._resolve_resume_style(_cfg(style, accent="default"))
    picked = rg._resolve_resume_style(_cfg(style, accent="burgundy"))
    assert picked["accent"] == default["accent"]
    assert picked["accent"] != rg.ACCENT_CHOICES["burgundy"]


def test_company_style_enum_survives_accent_substitution():
    """Regression: company_style's legit value is the string "accent" — the
    color-sentinel substitution must not rewrite it to a hex."""
    s = rg._resolve_resume_style(_cfg("ledger", accent="burgundy"))
    assert s["company_style"] == "accent"


def test_unknown_accent_falls_back_to_preset_accent():
    s = rg._resolve_resume_style(_cfg("ledger", accent="chartreuse"))
    assert s["accent"] == rg._resolve_resume_style(_cfg("ledger"))["accent"]


# ---------------------------------------------------------------------------
# The ATS invariant — single-column real text, nothing exotic
# ---------------------------------------------------------------------------
@pytest.mark.parametrize("style", STYLES)
def test_docx_has_no_ats_hostile_constructs(style, tmp_path, monkeypatch):
    monkeypatch.setattr(rg, "RESUMES_DIR", tmp_path)
    path = rg.generate_resume_docx(RESUME_CONTENT, PROFILE, JOB, _cfg(style))
    xml = _document_xml(path)

    assert "<w:tbl>" not in xml, f"{style}: tables break ATS parsing"
    assert "<w:drawing>" not in xml and "<w:pict>" not in xml, f"{style}: no images"
    assert "<w:txbxContent>" not in xml, f"{style}: no text boxes"
    # No multi-column section layout.
    assert 'w:num="2"' not in xml and 'w:num="3"' not in xml


@pytest.mark.parametrize("style", STYLES)
def test_docx_text_survives_extraction(style, tmp_path, monkeypatch):
    """What an ATS actually reads back out of the file."""
    monkeypatch.setattr(rg, "RESUMES_DIR", tmp_path)
    path = rg.generate_resume_docx(RESUME_CONTENT, PROFILE, JOB, _cfg(style))
    text = "\n".join(p.text for p in Document(path).paragraphs)

    for field in ("Morgan Reyes", "morgan.reyes@example.com", "Senior Data Analyst",
                  "Northwind Logistics", "2021 - Present", "Snowflake",
                  "B.S. Statistics", "Tableau Desktop Certified Professional",
                  "Dana Whitfield"):
        assert field in text, f"{style}: {field!r} lost in extraction"


def test_banner_band_is_paragraph_shading_not_a_table(tmp_path, monkeypatch):
    """The Banner band must be w:shd on real paragraphs — a table or image
    would defeat the whole point."""
    monkeypatch.setattr(rg, "RESUMES_DIR", tmp_path)
    path = rg.generate_resume_docx(RESUME_CONTENT, PROFILE, JOB, _cfg("banner"))
    xml = _document_xml(path)
    assert "<w:shd " in xml
    assert "<w:tbl>" not in xml


def test_executive_uses_small_caps_and_a_double_rule(tmp_path, monkeypatch):
    monkeypatch.setattr(rg, "RESUMES_DIR", tmp_path)
    path = rg.generate_resume_docx(RESUME_CONTENT, PROFILE, JOB, _cfg("executive"))
    xml = _document_xml(path)
    assert "<w:smallCaps/>" in xml
    assert 'w:val="double"' in xml


@pytest.mark.parametrize("style", STYLES)
def test_section_headers_keep_with_next(style, tmp_path, monkeypatch):
    """Headers must never be orphaned at the foot of a page."""
    monkeypatch.setattr(rg, "RESUMES_DIR", tmp_path)
    path = rg.generate_resume_docx(RESUME_CONTENT, PROFILE, JOB, _cfg(style))
    assert "<w:keepNext/>" in _document_xml(path)


# ---------------------------------------------------------------------------
# Cover letters follow the resume's style
# ---------------------------------------------------------------------------
@pytest.mark.parametrize("style", STYLES)
def test_cover_letter_renders_in_the_resume_style(style, tmp_path, monkeypatch):
    monkeypatch.setattr(rg, "RESUMES_DIR", tmp_path)
    path = rg.generate_cover_letter_docx(LETTER_CONTENT, PROFILE, JOB, _cfg(style))
    xml = _document_xml(path)
    s = rg._resolve_resume_style(_cfg(style))

    assert s["body_font"] in xml, f"{style}: cover letter ignored the style's font"
    assert "<w:tbl>" not in xml
    text = "\n".join(p.text for p in Document(path).paragraphs)
    assert "Morgan Reyes" in text
    assert "Meridian Freight" in text


def test_cover_letter_banner_matches_resume_banner(tmp_path, monkeypatch):
    monkeypatch.setattr(rg, "RESUMES_DIR", tmp_path)
    path = rg.generate_cover_letter_docx(LETTER_CONTENT, PROFILE, JOB, _cfg("banner"))
    assert "<w:shd " in _document_xml(path)


def test_cover_letter_drops_ai_written_closing(tmp_path, monkeypatch):
    """The renderer adds its own "Sincerely," — the AI's must not double up."""
    monkeypatch.setattr(rg, "RESUMES_DIR", tmp_path)
    path = rg.generate_cover_letter_docx(LETTER_CONTENT, PROFILE, JOB, _cfg())
    text = "\n".join(p.text for p in Document(path).paragraphs)
    assert text.count("Sincerely,") == 1


# ---------------------------------------------------------------------------
# PDF path
# ---------------------------------------------------------------------------
@pytest.mark.parametrize("style", STYLES)
def test_pdf_renders_for_every_style(style, tmp_path, monkeypatch):
    monkeypatch.setattr(rg, "RESUMES_DIR", tmp_path)
    path = rg._render_resume_pdf(RESUME_CONTENT, PROFILE, JOB, _cfg(style, fmt="pdf"))
    assert (tmp_path / "t1_resume.pdf").exists()
    with open(path, "rb") as fh:
        assert fh.read(5) == b"%PDF-"


def test_pdf_bullet_downgrades_unrenderable_glyphs():
    assert rg._pdf_bullet("▸  ") == "•  "   # decorative → safe round bullet
    assert rg._pdf_bullet("—  ") == "—  "   # em dash is WinAnsi-safe (Swiss)
    assert rg._pdf_bullet("-  ") == "-  "   # ASCII passes through


def test_swiss_pdf_stays_on_the_builtin_grotesque():
    """Helvetica *is* the Swiss design — it must not get swapped for Lato."""
    assert rg._pdf_font("Arial") == "Helvetica"
    assert rg._pdf_font("Arial", bold=True) == "Helvetica-Bold"


def test_dispatcher_passes_style_through_to_the_cover_letter(tmp_path, monkeypatch):
    """Regression: the cover letter used to ignore the style preset entirely."""
    monkeypatch.setattr(rg, "RESUMES_DIR", tmp_path)
    rg.generate_cover_letter(LETTER_CONTENT, PROFILE, JOB, _cfg("executive"))
    xml = _document_xml(tmp_path / "t1_cover_letter.docx")
    assert "Georgia" in xml  # executive's serif reached the letter


# ---------------------------------------------------------------------------
# Style preview (the settings picker)
# ---------------------------------------------------------------------------
@pytest.mark.parametrize("style", STYLES)
def test_preview_renders_every_style_to_a_pdf(style):
    pdf = rg.render_style_preview(style)
    assert pdf.startswith(b"%PDF-")
    assert len(pdf) > 1000        # a real page, not an empty document


@pytest.mark.parametrize("style", STYLES)
def test_preview_writes_nothing_to_disk(style, tmp_path, monkeypatch):
    """The preview is a look, not a document — it must not touch user files."""
    monkeypatch.setattr(rg, "RESUMES_DIR", tmp_path)
    rg.render_style_preview(style, "burgundy")
    assert list(tmp_path.iterdir()) == []


def test_preview_honors_accent_and_the_monochrome_lock():
    """The picker must not imply an accent applies where the style ignores it."""
    def accent_of(style, accent):
        cfg = {"application_honesty": {
            "resume_style": style, "resume_accent": accent,
        }}
        return rg._resolve_resume_style(cfg)["accent_color"]

    # Accent-bearing styles take the user's choice...
    assert accent_of("ledger", "burgundy") == rg.ACCENT_CHOICES["burgundy"]
    assert accent_of("banner", "forest") == rg.ACCENT_CHOICES["forest"]
    # ...and the accent_locked ones ignore it entirely.
    for style in ("executive", "swiss"):
        assert accent_of(style, "burgundy") == accent_of(style, "default")


def test_preview_sample_exercises_every_section_the_styles_render():
    """A specimen that skipped a section would hide that section's styling."""
    sections = rg._parse_resume_sections(rg.PREVIEW_CONTENT)
    assert {"summary", "skills", "experience", "education", "certifications"} <= set(sections)
    entries = rg._parse_experience_entries(sections["experience"])
    assert len(entries) >= 2                      # entry spacing is visible
    assert all(e["bullets"] for e in entries)     # bullet markers are visible
    assert all(e["dates"] for e in entries)       # right-aligned dates are visible
