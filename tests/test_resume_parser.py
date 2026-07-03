"""
Tests for backend/resume_parser.py — text extraction and the strictly-
extractive LLM mapping. All offline; the LLM client is mocked.
"""

import io
import json
from types import SimpleNamespace
from unittest.mock import AsyncMock, patch

import pytest

from backend import resume_parser


# ---------------------------------------------------------------------------
# extract_text
# ---------------------------------------------------------------------------
def test_extract_text_plain():
    assert resume_parser.extract_text("resume.txt", b"Hello world\n") == "Hello world"


def test_extract_text_unknown_extension_raises():
    with pytest.raises(ValueError):
        resume_parser.extract_text("resume.rtf", b"junk")


def test_extract_text_docx_roundtrip(tmp_path):
    from docx import Document
    doc = Document()
    doc.add_paragraph("Jane Q. Public")
    doc.add_paragraph("Senior Engineer")
    out = tmp_path / "r.docx"
    doc.save(out)
    text = resume_parser.extract_text("r.docx", out.read_bytes())
    assert "Jane Q. Public" in text
    assert "Senior Engineer" in text


def test_extract_text_pdf_corrupt_raises():
    # Garbage bytes for a .pdf path should surface a clean ValueError,
    # not a raw pypdf exception.
    with pytest.raises(ValueError):
        resume_parser.extract_text("r.pdf", b"not really a pdf")


# ---------------------------------------------------------------------------
# parse_resume — LLM mocked
# ---------------------------------------------------------------------------
def _mock_client(content: str):
    """Return a fake AsyncOpenAI-like client whose chat call yields `content`."""
    response = SimpleNamespace(
        choices=[SimpleNamespace(message=SimpleNamespace(content=content))]
    )
    client = SimpleNamespace()
    client.chat = SimpleNamespace()
    client.chat.completions = SimpleNamespace(create=AsyncMock(return_value=response))
    return client


@pytest.fixture
def cfg():
    return {"ai": {"base_url": "http://x", "models": {"strong": {"model": "m"}}}}


@pytest.mark.asyncio
async def test_parse_resume_empty_text(cfg):
    out = await resume_parser.parse_resume("", cfg)
    assert out["profile"]["full_name"] == ""
    assert out["profile"]["experience"] == []
    assert out["warnings"]


@pytest.mark.asyncio
async def test_parse_resume_extracts_clean_json(cfg):
    payload = {
        "full_name": "Ada Lovelace",
        "email": "ada@example.com",
        "phone": "555-1234",
        "location": "London, UK",
        "linkedin": "https://linkedin.com/in/ada",
        "summary": "Pioneering programmer.",
        "skills": ["Analytical Engine", "Mathematics"],
        "experience": [
            {"title": "Translator", "company": "Babbage Lab",
             "start_date": "1842", "end_date": "1843",
             "bullets": ["Wrote first algorithm"]},
        ],
        "education": [{"degree": "Tutoring", "school": "Home", "year": "1830"}],
        "certifications": ["RSS Fellow"],
    }
    client = _mock_client(json.dumps(payload))
    with patch.object(resume_parser.ai_engine, "_get_client", return_value=client), \
         patch.object(resume_parser.ai_engine, "_model", return_value="m"):
        out = await resume_parser.parse_resume("some resume text", cfg)

    p = out["profile"]
    assert p["full_name"] == "Ada Lovelace"
    assert p["email"] == "ada@example.com"
    assert p["skills"] == ["Analytical Engine", "Mathematics"]
    assert p["experience"][0]["title"] == "Translator"
    assert p["experience"][0]["bullets"] == ["Wrote first algorithm"]
    assert p["education"][0]["school"] == "Home"
    assert p["certifications"] == ["RSS Fellow"]


@pytest.mark.asyncio
async def test_parse_resume_strips_markdown_fence(cfg):
    fenced = "```json\n" + json.dumps({"full_name": "X", "email": "x@y.z"}) + "\n```"
    client = _mock_client(fenced)
    with patch.object(resume_parser.ai_engine, "_get_client", return_value=client), \
         patch.object(resume_parser.ai_engine, "_model", return_value="m"):
        out = await resume_parser.parse_resume("text", cfg)
    assert out["profile"]["full_name"] == "X"


@pytest.mark.asyncio
async def test_parse_resume_no_fabrication_on_empty_payload(cfg):
    # An honest model facing thin input may return all-empty values — these
    # MUST survive to the UI as empties, never as invented strings.
    client = _mock_client(json.dumps({
        "full_name": "", "email": "", "phone": "", "summary": "",
        "skills": [], "experience": [], "education": [], "certifications": [],
    }))
    with patch.object(resume_parser.ai_engine, "_get_client", return_value=client), \
         patch.object(resume_parser.ai_engine, "_model", return_value="m"):
        out = await resume_parser.parse_resume("a tiny resume", cfg)

    p = out["profile"]
    assert p["full_name"] == ""
    assert p["email"] == ""
    assert p["experience"] == []
    assert p["education"] == []
    assert p["skills"] == []


@pytest.mark.asyncio
async def test_parse_resume_garbage_model_output_safe(cfg):
    client = _mock_client("not json at all, sorry")
    with patch.object(resume_parser.ai_engine, "_get_client", return_value=client), \
         patch.object(resume_parser.ai_engine, "_model", return_value="m"):
        out = await resume_parser.parse_resume("text", cfg)
    assert out["profile"]["full_name"] == ""
    assert out["warnings"]


@pytest.mark.asyncio
async def test_parse_resume_drops_empty_experience_rows(cfg):
    client = _mock_client(json.dumps({
        "full_name": "Tester",
        "experience": [
            {"title": "Engineer", "company": "Acme",
             "start_date": "2020", "end_date": "Present", "bullets": []},
            {"title": "", "company": "", "bullets": []},
        ],
    }))
    with patch.object(resume_parser.ai_engine, "_get_client", return_value=client), \
         patch.object(resume_parser.ai_engine, "_model", return_value="m"):
        out = await resume_parser.parse_resume("text", cfg)
    exp = out["profile"]["experience"]
    assert len(exp) == 1
    assert exp[0]["company"] == "Acme"


@pytest.mark.asyncio
async def test_parse_resume_llm_failure_returns_warning(cfg):
    client = SimpleNamespace()
    client.chat = SimpleNamespace()
    client.chat.completions = SimpleNamespace(
        create=AsyncMock(side_effect=RuntimeError("LM Studio down"))
    )
    with patch.object(resume_parser.ai_engine, "_get_client", return_value=client), \
         patch.object(resume_parser.ai_engine, "_model", return_value="m"):
        out = await resume_parser.parse_resume("text", cfg)
    assert out["profile"]["full_name"] == ""
    assert any("AI extraction failed" in w for w in out["warnings"])
